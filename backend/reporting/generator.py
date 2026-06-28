"""Report generation in multiple formats: Markdown, HTML, JSON, CSV, Excel, PDF, DOCX."""

import io
import json
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from backend.core.logger import logger

try:
    import pandas as pd
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    _EXCEL_AVAILABLE = True
except ImportError:
    _EXCEL_AVAILABLE = False
    logger.warning("openpyxl/pandas not installed — Excel export disabled")

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        SimpleDocTemplate, Table, TableStyle, Paragraph,
        Spacer, HRFlowable, PageBreak,
    )
    _PDF_AVAILABLE = True
except ImportError:
    _PDF_AVAILABLE = False
    logger.warning("reportlab not installed — PDF export disabled")

try:
    from docx import Document as DocxDoc
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False
    logger.warning("python-docx not installed — DOCX export disabled")

try:
    from jinja2 import Template
    _JINJA_AVAILABLE = True
except ImportError:
    _JINJA_AVAILABLE = False


# ─── Report Generator ─────────────────────────────────────────────────────────

class ReportGenerator:
    """Generate test execution reports in multiple formats."""

    def __init__(self, report):
        """Accept an ExecutionReport or a plain dict."""
        if hasattr(report, "to_dict"):
            self._data = report.to_dict()
            self._report = report
        else:
            self._data = report
            self._report = None

        self.timestamp = datetime.utcnow()
        self.run_id = self._data.get("run_id", "unknown")
        self.results = self._data.get("results", [])
        self.summary = {
            "total_tests": self._data.get("total_tests", len(self.results)),
            "passed": self._data.get("passed", sum(1 for r in self.results if r.get("status") == "passed")),
            "failed": self._data.get("failed", sum(1 for r in self.results if r.get("status") == "failed")),
            "skipped": self._data.get("skipped", sum(1 for r in self.results if r.get("status") == "skipped")),
            "success_rate": self._data.get("success_rate", 0.0),
            "failure_rate": self._data.get("failure_rate", 0.0),
            "duration": self._data.get("duration", 0.0),
            "framework": self._data.get("framework", "generic"),
        }

    # ─── Markdown ─────────────────────────────────────────────────────────────

    def generate_markdown(self) -> str:
        s = self.summary
        md = f"""# QA Copilot — Test Execution Report

**Run ID:** `{self.run_id}`
**Framework:** {s['framework']}
**Generated:** {self.timestamp.strftime('%Y-%m-%d %H:%M:%S UTC')}

---

## Summary

| Metric | Value |
|--------|-------|
| Total Tests | {s['total_tests']} |
| ✅ Passed | {s['passed']} |
| ❌ Failed | {s['failed']} |
| ⏭ Skipped | {s['skipped']} |
| Success Rate | {s['success_rate']:.2f}% |
| Duration | {s['duration']:.2f}s |

---

## Test Results

"""
        for r in self.results:
            icon = {"passed": "✅", "failed": "❌", "skipped": "⏭", "error": "💥"}.get(r.get("status", ""), "❓")
            md += f"### {icon} {r.get('test_id', '?')} — {r.get('test_name', 'Unknown')}\n"
            md += f"- **Status:** {r.get('status', 'unknown')}\n"
            md += f"- **Duration:** {r.get('duration', 0):.3f}s\n"
            if r.get("error_message"):
                md += f"- **Error:** {r['error_message']}\n"
            if r.get("stack_trace"):
                md += f"\n```\n{r['stack_trace']}\n```\n"
            md += "\n"

        return md

    # ─── JSON ─────────────────────────────────────────────────────────────────

    def generate_json(self) -> str:
        output = {
            "run_id": self.run_id,
            "generated_at": self.timestamp.isoformat(),
            "summary": self.summary,
            "results": self.results,
        }
        return json.dumps(output, indent=2, default=str)

    # ─── HTML ─────────────────────────────────────────────────────────────────

    def generate_html(self) -> str:
        s = self.summary
        status_color = "#22c55e" if s["success_rate"] >= 80 else "#f59e0b" if s["success_rate"] >= 50 else "#ef4444"

        rows = ""
        for r in self.results:
            status = r.get("status", "unknown")
            color = {"passed": "#22c55e", "failed": "#ef4444", "skipped": "#f59e0b", "error": "#dc2626"}.get(status, "#6b7280")
            err = r.get("error_message", "") or ""
            rows += f"""<tr>
            <td>{r.get('test_id', '?')}</td>
            <td>{r.get('test_name', 'Unknown')}</td>
            <td style="color:{color};font-weight:bold">{status.upper()}</td>
            <td>{r.get('duration', 0):.3f}s</td>
            <td title="{err}">{err[:80] + '...' if len(err) > 80 else err}</td>
        </tr>"""

        return f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>QA Copilot Report — {self.run_id}</title>
    <style>
        * {{ box-sizing: border-box; margin: 0; padding: 0; }}
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
               background: #f8fafc; color: #1e293b; padding: 2rem; }}
        h1 {{ font-size: 1.875rem; font-weight: 700; margin-bottom: 0.5rem; }}
        .meta {{ color: #64748b; font-size: 0.875rem; margin-bottom: 2rem; }}
        .cards {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(150px, 1fr));
                  gap: 1rem; margin-bottom: 2rem; }}
        .card {{ background: white; border-radius: 0.75rem; padding: 1.25rem;
                 box-shadow: 0 1px 3px rgba(0,0,0,.1); }}
        .card .value {{ font-size: 2rem; font-weight: 700; color: {status_color}; }}
        .card .label {{ font-size: 0.75rem; color: #64748b; text-transform: uppercase;
                        letter-spacing: .05em; margin-top: 0.25rem; }}
        table {{ width: 100%; border-collapse: collapse; background: white;
                 border-radius: 0.75rem; overflow: hidden;
                 box-shadow: 0 1px 3px rgba(0,0,0,.1); }}
        th {{ background: #1e293b; color: white; padding: 0.75rem 1rem;
              text-align: left; font-size: 0.75rem; text-transform: uppercase; }}
        td {{ padding: 0.75rem 1rem; border-bottom: 1px solid #e2e8f0; font-size: 0.875rem; }}
        tr:hover td {{ background: #f1f5f9; }}
        tr:last-child td {{ border-bottom: none; }}
    </style>
</head>
<body>
    <h1>QA Copilot Test Execution Report</h1>
    <div class="meta">Run ID: {self.run_id} &nbsp;|&nbsp; Framework: {s['framework']} &nbsp;|&nbsp;
        Generated: {self.timestamp.strftime('%Y-%m-%d %H:%M:%S UTC')}</div>

    <div class="cards">
        <div class="card"><div class="value">{s['total_tests']}</div><div class="label">Total Tests</div></div>
        <div class="card"><div class="value" style="color:#22c55e">{s['passed']}</div><div class="label">Passed</div></div>
        <div class="card"><div class="value" style="color:#ef4444">{s['failed']}</div><div class="label">Failed</div></div>
        <div class="card"><div class="value" style="color:#f59e0b">{s['skipped']}</div><div class="label">Skipped</div></div>
        <div class="card"><div class="value" style="color:{status_color}">{s['success_rate']:.1f}%</div><div class="label">Success Rate</div></div>
        <div class="card"><div class="value">{s['duration']:.2f}s</div><div class="label">Duration</div></div>
    </div>

    <table>
        <thead><tr>
            <th>Test ID</th><th>Test Name</th><th>Status</th><th>Duration</th><th>Error</th>
        </tr></thead>
        <tbody>{rows}</tbody>
    </table>
</body>
</html>"""

    # ─── CSV ──────────────────────────────────────────────────────────────────

    def generate_csv(self) -> str:
        lines = ["Test ID,Test Name,Status,Duration (s),Error Message"]
        for r in self.results:
            err = (r.get("error_message") or "").replace('"', '""')
            lines.append(
                f'{r.get("test_id","")},"{r.get("test_name","")}",'
                f'{r.get("status","")},{r.get("duration",0):.3f},"{err}"'
            )
        return "\n".join(lines)

    # ─── Excel ────────────────────────────────────────────────────────────────

    def generate_excel(self) -> bytes:
        if not _EXCEL_AVAILABLE:
            raise RuntimeError("openpyxl not installed — Excel export unavailable")

        wb = openpyxl.Workbook()

        # Summary sheet
        ws_sum = wb.active
        ws_sum.title = "Summary"
        s = self.summary

        header_fill = PatternFill(start_color="1E293B", end_color="1E293B", fill_type="solid")
        header_font = Font(color="FFFFFF", bold=True, size=12)

        ws_sum["A1"] = "QA Copilot — Test Execution Report"
        ws_sum["A1"].font = Font(bold=True, size=16)
        ws_sum["A2"] = f"Run ID: {self.run_id} | Generated: {self.timestamp.strftime('%Y-%m-%d %H:%M:%S')}"
        ws_sum["A2"].font = Font(italic=True, color="64748B")

        ws_sum.append([])
        ws_sum.append(["Metric", "Value"])
        for cell in ws_sum[ws_sum.max_row]:
            cell.fill = header_fill
            cell.font = header_font

        metrics = [
            ("Total Tests", s["total_tests"]),
            ("Passed", s["passed"]),
            ("Failed", s["failed"]),
            ("Skipped", s["skipped"]),
            ("Success Rate (%)", f"{s['success_rate']:.2f}"),
            ("Duration (s)", f"{s['duration']:.2f}"),
            ("Framework", s["framework"]),
        ]
        for metric, value in metrics:
            ws_sum.append([metric, value])

        ws_sum.column_dimensions["A"].width = 22
        ws_sum.column_dimensions["B"].width = 20

        # Results sheet
        ws_res = wb.create_sheet("Test Results")
        headers = ["Test ID", "Test Name", "Status", "Duration (s)", "Error Message", "Timestamp"]
        ws_res.append(headers)
        for cell in ws_res[1]:
            cell.fill = header_fill
            cell.font = header_font

        status_colors = {
            "passed": "22C55E", "failed": "EF4444",
            "skipped": "F59E0B", "error": "DC2626",
        }
        for r in self.results:
            row = [
                r.get("test_id", ""),
                r.get("test_name", ""),
                r.get("status", ""),
                round(r.get("duration", 0), 3),
                r.get("error_message") or "",
                r.get("timestamp", ""),
            ]
            ws_res.append(row)
            status_cell = ws_res.cell(ws_res.max_row, 3)
            color = status_colors.get(r.get("status", ""), "6B7280")
            status_cell.fill = PatternFill(start_color=color, end_color=color, fill_type="solid")
            status_cell.font = Font(color="FFFFFF", bold=True)

        for col in range(1, len(headers) + 1):
            ws_res.column_dimensions[get_column_letter(col)].width = 20

        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()

    # ─── PDF ──────────────────────────────────────────────────────────────────

    def generate_pdf(self) -> bytes:
        if not _PDF_AVAILABLE:
            raise RuntimeError("reportlab not installed — PDF export unavailable")

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=0.75 * inch)
        styles = getSampleStyleSheet()
        story = []

        # Title
        title_style = ParagraphStyle("title", parent=styles["Heading1"], fontSize=20, spaceAfter=6)
        story.append(Paragraph("QA Copilot — Test Execution Report", title_style))
        story.append(Paragraph(
            f"Run ID: {self.run_id} | Generated: {self.timestamp.strftime('%Y-%m-%d %H:%M:%S UTC')}",
            styles["Normal"]
        ))
        story.append(Spacer(1, 0.25 * inch))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#1E293B")))
        story.append(Spacer(1, 0.25 * inch))

        # Summary table
        story.append(Paragraph("Summary", styles["Heading2"]))
        s = self.summary
        sum_data = [
            ["Metric", "Value"],
            ["Total Tests", str(s["total_tests"])],
            ["Passed", str(s["passed"])],
            ["Failed", str(s["failed"])],
            ["Skipped", str(s["skipped"])],
            ["Success Rate", f"{s['success_rate']:.2f}%"],
            ["Duration", f"{s['duration']:.2f}s"],
            ["Framework", s["framework"]],
        ]
        sum_table = Table(sum_data, colWidths=[3 * inch, 3 * inch])
        sum_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E293B")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F1F5F9")]),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
            ("ALIGN", (1, 0), (-1, -1), "CENTER"),
        ]))
        story.append(sum_table)
        story.append(Spacer(1, 0.5 * inch))

        # Results table
        story.append(Paragraph("Test Results", styles["Heading2"]))
        res_data = [["Test ID", "Test Name", "Status", "Duration", "Error"]]
        for r in self.results:
            err = (r.get("error_message") or "")[:60]
            res_data.append([
                r.get("test_id", ""),
                r.get("test_name", "")[:40],
                r.get("status", "").upper(),
                f"{r.get('duration', 0):.3f}s",
                err,
            ])

        status_fill_map = {
            "PASSED": colors.HexColor("#DCFCE7"),
            "FAILED": colors.HexColor("#FEE2E2"),
            "SKIPPED": colors.HexColor("#FEF3C7"),
            "ERROR": colors.HexColor("#FEE2E2"),
        }
        res_table = Table(res_data, colWidths=[1.2 * inch, 2 * inch, 0.8 * inch, 0.8 * inch, 2.4 * inch])
        style_cmds = [
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E293B")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CBD5E1")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8FAFC")]),
        ]
        for i, r in enumerate(self.results, start=1):
            status = r.get("status", "").upper()
            fill = status_fill_map.get(status)
            if fill:
                style_cmds.append(("BACKGROUND", (2, i), (2, i), fill))
        res_table.setStyle(TableStyle(style_cmds))
        story.append(res_table)

        doc.build(story)
        return buf.getvalue()

    # ─── DOCX ─────────────────────────────────────────────────────────────────

    def generate_docx(self) -> bytes:
        if not _DOCX_AVAILABLE:
            raise RuntimeError("python-docx not installed — DOCX export unavailable")

        doc = DocxDoc()
        s = self.summary

        # Title
        title = doc.add_heading("QA Copilot — Test Execution Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(
            f"Run ID: {self.run_id} | Framework: {s['framework']} | "
            f"Generated: {self.timestamp.strftime('%Y-%m-%d %H:%M:%S UTC')}"
        )

        doc.add_heading("Summary", 1)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Metric"
        hdr[1].text = "Value"

        metrics = [
            ("Total Tests", str(s["total_tests"])),
            ("Passed", str(s["passed"])),
            ("Failed", str(s["failed"])),
            ("Skipped", str(s["skipped"])),
            ("Success Rate", f"{s['success_rate']:.2f}%"),
            ("Duration", f"{s['duration']:.2f}s"),
        ]
        for metric, value in metrics:
            row = table.add_row().cells
            row[0].text = metric
            row[1].text = value

        doc.add_paragraph()
        doc.add_heading("Test Results", 1)

        results_table = doc.add_table(rows=1, cols=5)
        results_table.style = "Table Grid"
        headers = ["Test ID", "Test Name", "Status", "Duration", "Error"]
        for i, h in enumerate(headers):
            results_table.rows[0].cells[i].text = h

        for r in self.results:
            row = results_table.add_row().cells
            row[0].text = r.get("test_id", "")
            row[1].text = r.get("test_name", "")
            row[2].text = r.get("status", "").upper()
            row[3].text = f"{r.get('duration', 0):.3f}s"
            row[4].text = (r.get("error_message") or "")[:100]

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ─── Confluence Markdown ───────────────────────────────────────────────────

    def generate_confluence(self) -> str:
        """Confluence-compatible wiki markup."""
        s = self.summary
        wiki = f"""h1. QA Copilot Test Execution Report

*Run ID:* {self.run_id}
*Framework:* {s['framework']}
*Generated:* {self.timestamp.strftime('%Y-%m-%d %H:%M:%S UTC')}

----

h2. Summary

|| Metric || Value ||
| Total Tests | {s['total_tests']} |
| Passed | {{color:green}}{s['passed']}{{color}} |
| Failed | {{color:red}}{s['failed']}{{color}} |
| Skipped | {{color:orange}}{s['skipped']}{{color}} |
| Success Rate | {s['success_rate']:.2f}% |
| Duration | {s['duration']:.2f}s |

----

h2. Test Results

|| Test ID || Test Name || Status || Duration || Error ||
"""
        for r in self.results:
            status = r.get("status", "unknown")
            color = {"passed": "green", "failed": "red", "skipped": "orange"}.get(status, "gray")
            wiki += (
                f"| {r.get('test_id', '')} | {r.get('test_name', '')} | "
                f"{{color:{color}}}{status.upper()}{{color}} | "
                f"{r.get('duration', 0):.3f}s | {r.get('error_message') or '-'} |\n"
            )
        return wiki

    # ─── File Save ────────────────────────────────────────────────────────────

    def save_report(self, output_dir: str, fmt: str = "markdown") -> str:
        out = Path(output_dir)
        out.mkdir(parents=True, exist_ok=True)

        generators = {
            "markdown": (self.generate_markdown, f"report_{self.run_id}.md", "w"),
            "json": (self.generate_json, f"report_{self.run_id}.json", "w"),
            "html": (self.generate_html, f"report_{self.run_id}.html", "w"),
            "csv": (self.generate_csv, f"report_{self.run_id}.csv", "w"),
            "excel": (self.generate_excel, f"report_{self.run_id}.xlsx", "wb"),
            "pdf": (self.generate_pdf, f"report_{self.run_id}.pdf", "wb"),
            "docx": (self.generate_docx, f"report_{self.run_id}.docx", "wb"),
            "confluence": (self.generate_confluence, f"report_{self.run_id}_confluence.txt", "w"),
        }

        if fmt not in generators:
            raise ValueError(f"Unsupported format: {fmt}. Choose from: {list(generators)}")

        gen_fn, filename, mode = generators[fmt]
        content = gen_fn()
        filepath = out / filename

        if "b" in mode:
            filepath.write_bytes(content)
        else:
            filepath.write_text(content, encoding="utf-8")

        logger.info("Report saved: %s", filepath)
        return str(filepath)
